"""图片 → Word — 使用 OCR 识别图片中的文字并生成 Word 文档"""
import os
import threading
from typing import Optional

from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

from src.core.converter_base import BaseConverter
from src.core.converter_registry import register
from src.core.models import ConversionType
from src.utils.image_utils import preprocess_for_ocr


@register(ConversionType.IMAGE_TO_WORD)
class ImageToWordConverter(BaseConverter):
    """使用 OCR 将图片中的文字识别并导出为 Word"""

    def convert(
        self,
        input_path: str,
        output_dir: str,
        options: Optional[dict] = None,
        cancel_event: Optional[threading.Event] = None,
    ) -> str:
        options = options or {}
        lang = options.get("language", "chi_sim+eng")
        preprocess = options.get("preprocess", True)

        # 加载并预处理图片
        image = Image.open(input_path)
        if preprocess:
            image = preprocess_for_ocr(image)

        self._check_cancelled(cancel_event)

        # OCR 识别
        try:
            # 获取详细 OCR 数据
            ocr_data = pytesseract.image_to_data(
                image, lang=lang, output_type=pytesseract.Output.DICT
            )
            full_text = pytesseract.image_to_string(image, lang=lang)
        except pytesseract.TesseractNotFoundError:
            raise RuntimeError(
                "未找到 Tesseract OCR 引擎。请确保已安装 Tesseract OCR，"
                "或联系开发者获取包含 OCR 引擎的完整版本。"
            )

        self._check_cancelled(cancel_event)

        # 创建 Word 文档
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 检测是否包含表格结构
        has_table = self._detect_table_structure(ocr_data)

        if has_table:
            # 按表格结构写入
            self._write_as_table(doc, ocr_data)
        else:
            # 按段落写入
            for line in full_text.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph("")

        image.close()

        output_path = self._make_output_path(input_path, output_dir, ".docx")
        doc.save(output_path)
        return output_path

    def _detect_table_structure(self, ocr_data: dict) -> bool:
        """检测 OCR 结果中是否包含表格结构"""
        confidences = [
            int(c) for c in ocr_data.get("conf", [])
            if c != "-1" and int(c) > 30
        ]
        if len(confidences) < 10:
            return False

        # 检查列对齐情况（同一行中 x 坐标分布均匀则可能是表格）
        blocks = {}
        for i, (x, y, w, h, text) in enumerate(zip(
            ocr_data["left"], ocr_data["top"],
            ocr_data["width"], ocr_data["height"],
            ocr_data["text"]
        )):
            if text.strip():
                row_key = y // 20  # 按 20px 分组行
                if row_key not in blocks:
                    blocks[row_key] = []
                blocks[row_key].append(x)

        # 如果多行具有相似的 x 坐标分布，则可能是表格
        if len(blocks) >= 3:
            return True
        return False

    def _write_as_table(self, doc: Document, ocr_data: dict) -> None:
        """将 OCR 数据按表格形式写入 Word"""
        # 按行分组
        rows = {}
        for i in range(len(ocr_data["text"])):
            text = ocr_data["text"][i].strip()
            if not text:
                continue
            conf = int(ocr_data["conf"][i]) if ocr_data["conf"][i] != "-1" else 0
            if conf < 30:
                continue

            y = ocr_data["top"][i]
            x = ocr_data["left"][i]
            row_key = y // 15
            if row_key not in rows:
                rows[row_key] = []
            rows[row_key].append((x, text))

        if len(rows) < 2:
            # 数据不够表格化，回退为段落模式
            for row_key in sorted(rows.keys()):
                line = " ".join(t for _, t in sorted(rows[row_key]))
                doc.add_paragraph(line)
            return

        # 确定列数
        sorted_rows = sorted(rows.keys())
        max_cols = max(len(cells) for cells in rows.values())
        if max_cols < 2:
            for row_key in sorted_rows:
                line = " ".join(t for _, t in sorted(rows[row_key]))
                doc.add_paragraph(line)
            return

        table = doc.add_table(rows=len(sorted_rows), cols=max_cols, style="Table Grid")
        for i, row_key in enumerate(sorted_rows):
            cells = sorted(rows[row_key], key=lambda t: t[0])
            for j, (_, text) in enumerate(cells):
                if j < max_cols:
                    table.cell(i, j).text = text
